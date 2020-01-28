
import os
import re
from docx import Document
from google.cloud import translate_v2 as translate

# Create your own account in Google Cloud Platform and generate a file .json
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = 'my_token.json'
target = 'pl'  # you can change to another language
g_tr = translate.Client()  # utworzenie instancji klasy translatora


def doc_trans(fn):
    ''' Translate docx document '''
    doc = Document(fn)

    for para in doc.paragraphs:  # petla po paragrafach
        # jezeli paragraf nie jest pusty probuj tłumaczyc
        if para.text != "":
            try:
                tr_raw = g_tr.translate(para.text, target_language=target)
                tr = tr_raw['translatedText']
            except Exception as e:
                print(
                    ''' Google Translation error, 
                    check connection or your google .json''', e)
                return
            try:
                # nie ma różnicy pomiedzy tłumaczeniem a paragrafem
                if para.text == tr:
                    continue  # przejdź do następnego paragrafu
                # jeżeli znak specjalny na poczatku to nie tłumacz (zapewne kod)
                elif re.search(r'^[#!$<>=()@&\[\]{}_*]', para.text):
                    continue
                # jeżeli występuje znak '=' w paragrafie to tez nie tłumacz
                # to zapewne kod komputerowy
                elif re.search(r'=', para.text):
                    continue

                # Dodaj tłumaczenie z kursywa
                para.add_run('\n' + tr + '\n').italic = True

            except Exception as e:
                print('Error para.text', e)
            print(para.text)

    # save
    n_dir = os.path.dirname(fn)
    n_file = os.path.basename(fn)
    to_save = os.path.join(n_dir, target+'-'+n_file)
    doc.save(to_save)


if __name__ == '__main__':
    # from sys import argv
    # fn = argv[1]    # nazwa pliku docx do tłumaczenia
    doc_trans('test.docx')
